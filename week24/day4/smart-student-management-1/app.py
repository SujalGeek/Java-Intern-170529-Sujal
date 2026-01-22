import os
import json
import uuid
import re
from flask import Flask, request, jsonify
from langchain_community.document_loaders import PyPDFLoader
from langchain_groq import ChatGroq
from langchain_core.prompts import PromptTemplate
from langchain_classic.chains import LLMChain
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.colors import Color, gray
from docx import Document
from docx.shared import RGBColor
import numpy as np
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from sklearn.linear_model import LinearRegression

app = Flask(__name__)

# ==========================================
# CONFIGURATION
# ==========================================
os.environ["GROQ_API_KEY"] = ""
OUTPUT_FOLDER = "generated_exams"

if not os.path.exists(OUTPUT_FOLDER):
    os.makedirs(OUTPUT_FOLDER)

# ==========================================
# 1. AI LOGIC (Bloom's Taxonomy)
# ==========================================
def generate_exam_json(book_path):
    print(f"Reading book path: {book_path}")

    if not os.path.exists(book_path):
        raise FileNotFoundError(f"File not found at: {book_path}")

    file_size = os.path.getsize(book_path)
    if file_size == 0:
        raise ValueError("The uploaded file is 0 bytes (Empty). Upload failed.")

    loader = PyPDFLoader(book_path)
    pages = loader.load()

    context_text = ""
    start_page = 1 
    end_page = min(len(pages), 17) 

    for page in pages[start_page:end_page]:
        context_text += page.page_content + "\n"

    print("Asking AI to generate the Bloom's Taxonomy Exam...")

    prompt_text = """
    You are an Expert Professor. Generate a Mid-Term Exam based on the text below.

    REQUIREMENTS:
    1. Section A: 5 MCQs (Bloom's Levels: Remembering, Understanding).
    2. Section B: 4 Short Answer Questions (Bloom's Levels: Applying, Analyzing).

    TEXT:
    {text_content}

    OUTPUT FORMAT (Strict JSON):
    {{
      "exam_title": "Mid-Term Examination",
      "section_A_mcq": [
        {{ 
          "id": 1,
          "bloom": "Remembering",
          "question": "Question text here...",
          "options": ["A) Option 1", "B) Option 2", "C) Option 3", "D) Option 4"],
          "correct": "A",
          "explanation": "Brief explanation here..."
        }}
      ],
      "section_B_theory": [
        {{
          "id": 1,
          "bloom": "Applying",
          "question": "Question text here...",
          "model_answer": "Model answer here..."
        }}
      ]
    }}
    """

    llm = ChatGroq(
        model="llama-3.3-70b-versatile",
        temperature=0.3
    )

    prompt = PromptTemplate(
        template=prompt_text,
        input_variables=["text_content"]
    )

    chain = LLMChain(llm=llm, prompt=prompt)
    
    result = chain.run(context_text)

    print("--- RAW AI OUTPUT START ---")
    print(result)
    print("--- RAW AI OUTPUT END ---")

    try:
        json_match = re.search(r'\{.*\}', result, re.DOTALL)
        
        if json_match:
            clean_json = json_match.group(0)
            return json.loads(clean_json)
        else:
            clean_json = result.replace("```json", "").replace("```", "").strip()
            return json.loads(clean_json)
            
    except json.JSONDecodeError as e:
        print(f"JSON Parsing Failed: {e}")
        return {
            "exam_title": "Error Generating Exam",
            "section_A_mcq": [],
            "section_B_theory": []
        }

# ==========================================
# 2. PDF GENERATION (Student Copy)
# ==========================================
def create_watermarked_pdf(data, institute_name):
    filename = f"{OUTPUT_FOLDER}/Student_Copy_{uuid.uuid4().hex[:6]}.pdf"
    c = canvas.Canvas(filename, pagesize=letter)
    width, height = letter

    def draw_watermark():
        c.saveState()
        c.setFont("Helvetica-Bold", 60)
        c.setFillColor(Color(0.9, 0.9, 0.9))
        c.translate(width / 2, height / 2)
        c.rotate(45)
        c.drawCentredString(0, 0, institute_name)
        c.restoreState()

    y = height - 50
    draw_watermark()

    c.setFont("Helvetica-Bold", 16)
    c.drawCentredString(width / 2, y, data.get("exam_title", "Exam Paper"))
    y -= 30

    c.setFont("Helvetica", 10)
    c.drawCentredString(width / 2, y, f"Institute: {institute_name} | Time: 1 Hr")
    c.line(50, y - 10, width - 50, y - 10)
    y -= 50

    c.setFont("Helvetica-Bold", 12)
    c.drawString(50, y, "Section A: MCQs")
    y -= 20
    c.setFont("Helvetica", 10)

    for item in data.get("section_A_mcq", []):
        if y < 100:
            c.showPage()
            draw_watermark()
            y = height - 50

        c.drawString(50, y, f"Q{item['id']}. {item['question']}")
        y -= 15

        for opt in item.get("options", []):
            c.drawString(70, y, str(opt))
            y -= 12

        y -= 10

    y -= 20
    c.setFont("Helvetica-Bold", 12)
    c.drawString(50, y, "Section B: Theory")
    y -= 20
    c.setFont("Helvetica", 10)

    for item in data.get("section_B_theory", []):
        if y < 150:
            c.showPage()
            draw_watermark()
            y = height - 50

        c.drawString(50, y, f"Q{item['id']}. {item['question']}")
        c.setStrokeColor(gray)
        c.line(50, y - 20, width - 50, y - 20)
        c.line(50, y - 40, width - 50, y - 40)
        y -= 60

    c.save()
    return filename

# ==========================================
# 3. DOCX GENERATION (Teacher Key)
# ==========================================
def create_teacher_docx(data):
    filename = f"{OUTPUT_FOLDER}/Teacher_Key_{uuid.uuid4().hex[:6]}.docx"
    doc = Document()

    doc.add_heading(f"{data.get('exam_title', 'Exam')} - ANSWER KEY", 0)

    doc.add_heading("Section A: MCQs", level=1)

    for item in data.get("section_A_mcq", []):
        p = doc.add_paragraph()
        p.add_run(f"Q{item['id']}. {item['question']}").bold = True
        p.add_run(f" [Bloom: {item.get('bloom', 'General')}]").italic = True

        for opt in item.get("options", []):
            doc.add_paragraph(str(opt))

        ans = doc.add_paragraph(f"✅ Correct: {item.get('correct', 'N/A')}")
        if ans.runs:
            ans.runs[0].font.color.rgb = RGBColor(0, 128, 0)
        doc.add_paragraph(f"ℹ️ {item.get('explanation', '')}")

    doc.add_heading("Section B: Theory", level=1)

    for item in data.get("section_B_theory", []):
        p = doc.add_paragraph()
        p.add_run(f"Q{item['id']}. {item['question']}").bold = True
        doc.add_paragraph(f"Model Answer: {item.get('model_answer', '')}")

    doc.save(filename)
    return filename

# ==========================================
# 4. API ENDPOINTS
# ==========================================
@app.route("/generate-exam", methods=["POST"])
def generate_api():
    try:
        req = request.json
        file_path = req.get("filePath")
        institute = req.get("instituteName", "SMART STUDENT SYSTEM")

        exam_data = generate_exam_json(file_path)

        pdf_path = create_watermarked_pdf(exam_data, institute)
        docx_path = create_teacher_docx(exam_data)

        return jsonify({
            "success": True,
            "studentPdf": os.path.abspath(pdf_path),
            "teacherDocx": os.path.abspath(docx_path),
            "rawData": exam_data
        })

    except Exception as e:
        print(f"Error: {e}")
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500

@app.route('/grade-theory', methods=['POST'])
def grade_theory():
    try:
        data = request.json
        student_answer = data.get('student_answer', "")
        model_answer = data.get('model_answer', "")
        max_marks = data.get('max_marks', 5)
        
        if not student_answer or not model_answer:
            return jsonify({"success": True, "marks_awarded": 0, "similarity_percentage": 0})
            
        documents = [model_answer, student_answer]
        tfidf_vectorizer = TfidfVectorizer()
        tfidf_matrix = tfidf_vectorizer.fit_transform(documents)
        
        similarity_score = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]
        
        awarded_marks = 0
        if similarity_score > 0.85:
            awarded_marks = max_marks
        elif similarity_score > 0.5:
            awarded_marks = max_marks * similarity_score
        elif similarity_score > 0.2:
            awarded_marks = 1
        else:
            awarded_marks = 0
            
        return jsonify({
            "success": True,
            "similarity_percentage": round(similarity_score * 100, 2),
            "marks_awarded": round(awarded_marks, 1),
            "max_marks": max_marks,
            "feedback": "Good match" if similarity_score > 0.7 else "Content mismatch or missing keywords."
        })
        
    except Exception as e:
        return jsonify({
            "success": False,
            "message": str(e)
        }), 500


@app.route('/predict-score', methods=['POST'])
def predict_score():
    try:
        data = request.json
        scores = data.get('scores', [])
        
        if not scores:
            return jsonify({
                "success": False,
                "message": "No scores provided"
            }), 400
        
        # FIX 1: Corrected 'len(scores)' instead of 'len(score)'
        if len(scores) < 2:
            last_score = scores[0] if scores else 0
            return jsonify({
                "success": True,
                "predicted_score": last_score,
                "understanding_level": "Insufficient Data",
                "recommendation": "Complete more assessments to generate analysis."
            })
            
        # FIX 2: Un-indented this block so it actually runs!
        X = np.array(range(len(scores))).reshape(-1, 1)
        y = np.array(scores).reshape(-1, 1)
        
        model = LinearRegression()
        model.fit(X, y)
        
        next_exam_index = np.array([[len(scores)]])
        prediction = model.predict(next_exam_index)[0][0]
        prediction = round(prediction, 2)
        
        slope = model.coef_[0][0]
        
        understanding_level = ""
        recommendation = ""
        
        if prediction >= 8.0:
            understanding_level = "Deep Conceptual Mastery"
            recommendation = "Student is ready for advanced modules."
            
        elif prediction >= 6.0:
            understanding_level = "Good Functional Grasp"
            
            if slope < 0:
                recommendation = "Warning: Performance is slipping. Review recent topics."
            else:
                recommendation = "Consistent performance. Keep practicing"
                
        elif prediction >= 4.0:
            understanding_level = "Superficial Understanding"
            recommendation = "Student relies on memory rather than logic. Needs conceptual review."
        
        else:
            understanding_level = "Critical Knowledge Gap"
            recommendation = "URGENT: Fundamentals are unclear. Schedule remedial session."

        return jsonify({
            "success": True,
            "predicted_score": prediction,
            "trend": "Improving" if slope > 0 else "Declining",
            "understanding_level": understanding_level,
            "action_plan": recommendation
        })

    # FIX 3: Aligned except block correctly
    except Exception as e:
        return jsonify({
            "success": False,
            "message": str(e)
        }), 500
			
if __name__ == "__main__":
    app.run(port=5000)