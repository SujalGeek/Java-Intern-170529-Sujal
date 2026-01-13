import os
import json
import uuid
from flask import Flask, request, jsonify
from langchain_community.document_loaders import PyPDFLoader
from langchain_groq import ChatGroq
from langchain_core.prompts import PromptTemplate
from langchain_classic.chains import LLMChain
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.colors import Color, gray
from docx import Document
from docx.shared import RGBColor

app = Flask(__name__)

# ==========================================
# CONFIGURATION
# ==========================================
os.environ["GROQ_API_KEY"] = ""
OUTPUT_FOLDER = "generated_exams"

if not os.path.exists(OUTPUT_FOLDER):
    os.makedirs(OUTPUT_FOLDER)

# ==========================================
# 1. AI LOGIC (Bloom's Taxonomy)
# ==========================================
def generate_exam_json(book_path):

    if not os.path.exists(book_path):
        raise FileNotFoundError(f"File not found at: {book_path}")

    file_size = os.path.getsize(book_path)
    print(f"File Size: {file_size} bytes")

    if file_size == 0:
        raise ValueError("The uploaded file is 0 bytes (Empty). Upload failed.")

    print(f"Reading book: {book_path}")

    loader = PyPDFLoader(book_path)
    pages = loader.load()

    # Extract pages 10–20 to skip ToC / Intro
    context_text = ""
    end_page = min(len(pages), 20)

    for page in pages[10:end_page]:
        context_text += page.page_content + "\n"

    print("Asking AI to generate the Bloom's Taxonomy Exam...")

    prompt_text = """
    You are an Expert Professor. Generate a Mid-Term Exam based on the text below.

    REQUIREMENTS:
    1. Section A: 5 MCQs (Bloom's Levels: Remembering, Understanding).
    2. Section B: 4 Short Answer Questions (Bloom's Levels: Applying, Analyzing).

    TEXT:
    {text_content}

    OUTPUT FORMAT (Strict JSON):
    {{
      "exam_title": "Mid-Term Examination",
      "section_A_mcq": [
        {{ 
          "id": 1,
          "bloom": "Remembering",
          "question": "...",
          "options": ["A)", "B)", "C)", "D)"],
          "correct": "A",
          "explanation": "..."
        }}
      ],
      "section_B_theory": [
        {{
          "id": 1,
          "bloom": "Applying",
          "question": "...",
          "model_answer": "..."
        }}
      ]
    }}
    """

    llm = ChatGroq(
        model="llama-3.3-70b-versatile",
        temperature=0.3
    )

    prompt = PromptTemplate(
        template=prompt_text,
        input_variables=["text_content"]
    )

    chain = LLMChain(llm=llm, prompt=prompt)
    result = chain.run(context_text)

    # Remove Markdown formatting if added
    clean_json = result.replace("```json", "").replace("```", "").strip()
    return json.loads(clean_json)

# ==========================================
# 2. PDF GENERATION (Student Copy)
# ==========================================
def create_watermarked_pdf(data, institute_name):

    filename = f"{OUTPUT_FOLDER}/Student_Copy_{uuid.uuid4().hex[:6]}.pdf"
    c = canvas.Canvas(filename, pagesize=letter)
    width, height = letter

    def draw_watermark():
        c.saveState()
        c.setFont("Helvetica-Bold", 60)
        c.setFillColor(Color(0.9, 0.9, 0.9))
        c.translate(width / 2, height / 2)
        c.rotate(45)
        c.drawCentredString(0, 0, institute_name)
        c.restoreState()

    y = height - 50
    draw_watermark()

    # Title
    c.setFont("Helvetica-Bold", 16)
    c.drawCentredString(width / 2, y, data.get("exam_title", "Exam Paper"))
    y -= 30

    c.setFont("Helvetica", 10)
    c.drawCentredString(
        width / 2, y, f"Institute: {institute_name} | Time: 1 Hr"
    )
    c.line(50, y - 10, width - 50, y - 10)
    y -= 50

    # Section A – MCQs
    c.setFont("Helvetica-Bold", 12)
    c.drawString(50, y, "Section A: MCQs")
    y -= 20
    c.setFont("Helvetica", 10)

    for item in data.get("section_A_mcq", []):
        if y < 100:
            c.showPage()
            draw_watermark()
            y = height - 50

        c.drawString(50, y, f"Q{item['id']}. {item['question']}")
        y -= 15

        for opt in item["options"]:
            c.drawString(70, y, opt)
            y -= 12

        y -= 10

    # Section B – Theory
    y -= 20
    c.setFont("Helvetica-Bold", 12)
    c.drawString(50, y, "Section B: Theory")
    y -= 20
    c.setFont("Helvetica", 10)

    for item in data.get("section_B_theory", []):
        if y < 150:
            c.showPage()
            draw_watermark()
            y = height - 50

        c.drawString(50, y, f"Q{item['id']}. {item['question']}")
        c.setStrokeColor(gray)
        c.line(50, y - 20, width - 50, y - 20)
        c.line(50, y - 40, width - 50, y - 40)
        y -= 60

    c.save()
    return filename

# ==========================================
# 3. DOCX GENERATION (Teacher Key)
# ==========================================
def create_teacher_docx(data):

    filename = f"{OUTPUT_FOLDER}/Teacher_Key_{uuid.uuid4().hex[:6]}.docx"
    doc = Document()

    doc.add_heading(
        f"{data.get('exam_title', 'Exam')} - ANSWER KEY", 0
    )

    doc.add_heading("Section A: MCQs", level=1)

    for item in data.get("section_A_mcq", []):
        p = doc.add_paragraph()
        p.add_run(f"Q{item['id']}. {item['question']}").bold = True
        p.add_run(
            f" [Bloom: {item.get('bloom', 'General')}]"
        ).italic = True

        for opt in item["options"]:
            doc.add_paragraph(opt)

        ans = doc.add_paragraph(f"✅ Correct: {item['correct']}")
        ans.runs[0].font.color.rgb = RGBColor(0, 128, 0)
        doc.add_paragraph(f"ℹ️ {item['explanation']}")

    doc.add_heading("Section B: Theory", level=1)

    for item in data.get("section_B_theory", []):
        p = doc.add_paragraph()
        p.add_run(f"Q{item['id']}. {item['question']}").bold = True
        doc.add_paragraph(f"Model Answer: {item['model_answer']}")

    doc.save(filename)
    return filename

# ==========================================
# 4. API ENDPOINT
# ==========================================
@app.route("/generate-exam", methods=["POST"])
def generate_api():

    try:
        req = request.json
        file_path = req.get("filePath")
        institute = req.get("instituteName", "SMART STUDENT SYSTEM")

        # AI Processing
        exam_data = generate_exam_json(file_path)

        # File Generation
        pdf_path = create_watermarked_pdf(exam_data, institute)
        docx_path = create_teacher_docx(exam_data)

        return jsonify({
            "success": True,
            "studentPdf": os.path.abspath(pdf_path),
            "teacherDocx": os.path.abspath(docx_path),
            "rawData": exam_data
        })

    except Exception as e:
        print(f"Error: {e}")
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500

if __name__ == "__main__":
    app.run(port=5000)
